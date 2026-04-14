"""DocCommander table-row formatting tests."""
import asyncio
from pathlib import Path
from uuid import uuid4

from docx import Document

from core.doc_commander import DocCommander
from llm.base import BaseLLM


TEST_DATA_DIR = Path(__file__).parent / "test_data"


class FakeLLM(BaseLLM):
    def __init__(self, response: str):
        self.response = response

    async def chat(self, messages: list[dict], temperature: float = 0.1) -> str:
        return self.response


def test_parse_command_maps_row_request_to_table_row():
    commander = DocCommander.__new__(DocCommander)
    commander.llm = FakeLLM(
        '```json{"action":"format","target":"paragraph","params":{"index":1,"bold":true},"description":"将第二行加粗"}```'
    )

    result = asyncio.run(commander.parse_command("把第二行加粗"))

    assert result["target"] == "table_row"
    assert result["params"]["index"] == 1


def test_execute_can_bold_second_table_row():
    source = TEST_DATA_DIR / "table_format_case.docx"
    working = TEST_DATA_DIR / f"table_format_case_working_{uuid4().hex}.docx"
    working.write_bytes(source.read_bytes())

    try:
        commander = DocCommander.__new__(DocCommander)
        result = commander.execute(
            str(working),
            {"action": "format", "target": "table_row", "params": {"index": 1, "bold": True}},
        )

        assert result["success"] is True

        doc = Document(str(working))
        second_row = doc.tables[0].rows[1]
        assert any(run.bold for cell in second_row.cells for paragraph in cell.paragraphs for run in paragraph.runs)
        del doc
    finally:
        try:
            working.chmod(0o666)
            working.unlink(missing_ok=True)
        except PermissionError:
            pass
