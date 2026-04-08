"""生成测试文档和模板表格 - 用于系统功能验证"""
import sys
import os
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from pathlib import Path
from docx import Document as DocxDocument
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.utils import get_column_letter

OUTPUT_DIR = Path(__file__).parent / "test_data" / "documents"
TEMPLATE_DIR = Path(__file__).parent / "test_data" / "templates"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)

# ============================================================
# 测试文档数据（包含丰富的实体信息用于提取测试）
# ============================================================

DOCS_DATA = [
    # --- 4个 docx ---
    {
        "filename": "项目合同书.docx", "type": "docx",
        "title": "软件开发项目合同书",
        "content": [
            "甲方：北京智远科技有限公司\n地址：北京市海淀区中关村大街88号\n联系人：王建国\n电话：010-82567890\n邮箱：wangjg@zhiyuan-tech.com",
            "乙方：上海数联信息技术有限公司\n地址：上海市浦东新区张江高科技园区碧波路690号\n联系人：李明华\n电话：021-58901234\n邮箱：limh@shulian.cn",
            "合同编号：ZY-2026-SJ-0088\n签订日期：2026年3月15日\n合同金额：人民币壹佰贰拾万元整（¥1,200,000.00）",
            "项目名称：企业数据中台建设项目\n项目周期：2026年4月1日至2026年9月30日\n项目经理：张伟（甲方）、陈思远（乙方）",
            "付款方式：分三期支付。第一期：合同签订后7个工作日内支付30%，即360,000元；第二期：中期验收通过后支付40%，即480,000元；第三期：终验通过后支付30%，即360,000元。",
            "验收标准：系统功能完整性达到95%以上，性能指标满足需求文档要求，安全测试通过等保二级认证。",
        ],
    },
    {
        "filename": "年度工作报告.docx", "type": "docx",
        "title": "2025年度工作总结报告",
        "content": [
            "报告人：赵丽娟\n部门：市场营销部\n日期：2026年1月15日",
            "一、年度业绩概述\n2025年度，市场营销部在总经理刘德明的领导下，全年实现营业收入8,560万元，同比增长23.5%。新签客户127家，客户续约率达到91.2%。",
            "二、重点项目\n1. 与中国移动通信集团签署战略合作协议，合同金额2,800万元，项目负责人：孙浩然，联系电话：13912345678。\n2. 完成华为技术有限公司的数字化转型咨询项目，项目金额560万元。\n3. 中标国家电网信息化升级项目，金额1,200万元，预计2026年6月完成交付。",
            "三、团队建设\n部门现有员工45人，2025年新入职12人。核心成员：\n- 高级经理：周文博，负责大客户管理，邮箱：zhouwb@company.com\n- 区域总监：吴晓燕，负责华东区域，电话：13887654321\n- 市场分析师：林志强，负责行业研究，身份证号：310101199503121234",
        ],
    },
    {
        "filename": "会议纪要.docx", "type": "docx",
        "title": "第三季度产品规划会议纪要",
        "content": [
            "会议时间：2026年3月10日 14:00-17:30\n会议地点：深圳市南山区科技园南路18号 创新大厦A座15楼会议室\n主持人：黄志强（产品总监）",
            "参会人员：\n1. 黄志强 - 产品总监\n2. 马丽萍 - 技术总监，电话：13566778899\n3. 杨建军 - 运营总监\n4. 刘芳 - UI设计主管\n5. 陈伟 - 后端开发组长",
            "会议议题一：Q3产品路线图\n决议：确定在2026年7月1日前完成V3.0版本开发，预算150万元。由马丽萍负责技术方案评审，杨建军负责运营推广计划。",
            "会议议题二：用户反馈处理\n当前活跃用户数达到58万，月均增长率12%。用户满意度评分4.2/5.0。需要在2026年4月15日前完成TOP10问题修复。",
            "下次会议：2026年4月8日 14:00，地点不变。请各部门提前准备Q3详细执行计划。",
        ],
    },
    {
        "filename": "招标公告.docx", "type": "docx",
        "title": "智慧校园信息化建设项目招标公告",
        "content": [
            "招标编号：GDZB-2026-0315\n招标人：广东省教育厅\n地址：广州市越秀区东风东路723号\n联系人：何晓明\n电话：020-37627890\n邮箱：hexm@gdedu.gov.cn",
            "项目名称：广东省智慧校园信息化建设项目\n项目预算：人民币叁仟伍佰万元整（¥35,000,000.00）\n资金来源：省级财政拨款",
            "投标截止时间：2026年4月20日 17:00\n开标时间：2026年4月21日 09:30\n开标地点：广州市天河区天河路385号太古汇办公楼28层开标室",
            "投标人资格要求：\n1. 注册资本不低于5000万元\n2. 具有信息系统集成二级及以上资质\n3. 近三年内完成过单项合同金额不低于1000万元的教育信息化项目\n4. 项目经理须持有PMP证书",
            "联系方式：\n采购代理机构：广州华信招标有限公司\n地址：广州市天河区体育西路103号维多利广场A座2601\n联系人：钱丽华\n电话：020-38762345\n传真：020-38762346",
        ],
    },
    # --- 4个 txt ---
    {
        "filename": "员工通知.txt", "type": "txt",
        "content": [
            "关于2026年清明节放假安排的通知",
            "各部门：",
            "根据国务院办公厅通知精神，现将2026年清明节放假安排通知如下：",
            "放假时间：2026年4月4日（星期六）至4月6日（星期一），共3天。4月7日（星期二）正常上班。",
            "注意事项：\n1. 各部门负责人请于4月3日前完成值班安排表并报送行政部王小红处，联系电话：内线8023。\n2. 假期期间如有紧急事务，请联系行政部值班电话：010-65432100。\n3. 请各位同事注意出行安全，保持通讯畅通。",
            "特此通知。\n\n北京创新科技股份有限公司\n行政人事部\n2026年3月20日",
        ],
    },
    {
        "filename": "产品说明书.txt", "type": "txt",
        "content": [
            "DocFusion 智能文档处理系统 V2.0 产品说明书",
            "开发商：杭州数智科技有限公司\n地址：杭州市余杭区文一西路1218号\n技术支持热线：400-888-9999\n邮箱：support@shuzhi-tech.com",
            "产品概述：DocFusion是一款基于大语言模型的智能文档处理系统，支持文档解析、信息提取、智能填表等功能。",
            "系统要求：\n- 操作系统：Windows 10/11 64位\n- 内存：8GB及以上\n- 硬盘空间：2GB可用空间\n- Python 3.10+（如使用源码部署）",
            "授权信息：\n许可证编号：DF-ENT-2026-00156\n授权单位：南京信息工程大学\n有效期：2026年1月1日至2026年12月31日\n授权金额：98,000元/年",
        ],
    },
    {
        "filename": "调研报告摘要.txt", "type": "txt",
        "content": [
            "中国人工智能产业发展调研报告（2025）摘要",
            "编写单位：中国信息通信研究院\n发布日期：2026年2月28日\n报告负责人：张晓峰 博士\n联系邮箱：zhangxf@caict.ac.cn",
            "核心数据：\n- 2025年中国AI产业规模达到5,820亿元，同比增长31.2%\n- AI企业数量超过4,200家，从业人员约85万人\n- 北京、上海、深圳三地AI企业数量占全国总量的62%",
            "重点企业：百度、阿里巴巴、腾讯、华为、科大讯飞、商汤科技等\n投融资情况：2025年AI领域融资总额达到1,250亿元，其中大模型方向融资占比38%",
            "政策建议：\n1. 加大基础研究投入，建议年度投入不低于200亿元\n2. 完善AI人才培养体系，目标到2028年培养AI专业人才50万人\n3. 推动AI在医疗、教育、制造业等领域的深度应用",
        ],
    },
    {
        "filename": "客户反馈记录.txt", "type": "txt",
        "content": [
            "客户反馈记录 - 2026年3月",
            "记录人：陈小芳\n部门：客户服务部\n日期：2026年3月18日",
            "反馈1：\n客户：深圳前海金融科技有限公司\n联系人：郑伟明\n电话：13923456789\n邮箱：zhengwm@qhfintech.com\n问题：系统在处理超过50页的PDF文档时响应时间超过3分钟\n优先级：高\n处理状态：已转交技术部处理",
            "反馈2：\n客户：成都天府软件园管理有限公司\n联系人：罗小琴\n电话：028-85231234\n问题：希望增加对WPS格式文档的支持\n优先级：中\n处理状态：已纳入V3.0需求池",
            "反馈3：\n客户：武汉光谷信息技术有限公司\n联系人：胡建华\n电话：027-87654321\n邮箱：hujh@wh-optics.com\n问题：实体提取的准确率在法律文书场景下偏低，约75%\n优先级：高\n处理状态：已安排Prompt优化，预计3月25日完成",
        ],
    },
]

MD_DOCS = [
    {
        "filename": "技术方案.md", "type": "md",
        "content": """# 数据中台技术方案

## 项目信息
- 项目名称：企业数据中台建设
- 客户：中国银行股份有限公司
- 项目经理：刘洋，电话：13811223344，邮箱：liuyang@dataplatform.cn
- 技术负责人：赵明，电话：13922334455

## 技术架构
采用微服务架构，基于 Kubernetes 容器编排平台部署。
- 数据采集层：Kafka + Flink 实时流处理
- 数据存储层：MySQL + ClickHouse + MinIO
- 数据服务层：Spring Cloud + GraphQL

## 预算明细
| 项目 | 金额 |
|------|------|
| 硬件采购 | 450,000元 |
| 软件许可 | 280,000元 |
| 人力成本 | 680,000元 |
| 合计 | 1,410,000元 |

## 里程碑
1. 需求确认：2026年4月15日
2. 架构评审：2026年5月1日
3. 开发完成：2026年7月31日
4. 上线运行：2026年8月15日
""",
    },
    {
        "filename": "团队介绍.md", "type": "md",
        "content": """# 研发团队介绍

## 公司简介
深圳市云智科技有限公司，成立于2020年6月，注册资本1000万元。
地址：深圳市南山区粤海街道高新南一道6号深圳湾科技生态园9栋B座

## 核心团队

### 张明远 - CEO
- 清华大学计算机系博士
- 前腾讯高级技术总监
- 电话：13500001111
- 邮箱：zhangmy@yunzhi.ai

### 李婷 - CTO
- 北京大学人工智能硕士
- 10年AI研发经验
- 电话：13600002222
- 邮箱：liting@yunzhi.ai

### 王磊 - 产品VP
- 浙江大学MBA
- 前阿里巴巴产品专家
- 电话：13700003333

## 联系我们
- 官网：www.yunzhi.ai
- 客服热线：400-123-4567
- 地址：深圳市南山区粤海街道高新南一道6号
""",
    },
    {
        "filename": "需求文档.md", "type": "md",
        "content": """# 智慧医疗系统需求规格说明书

## 文档信息
- 版本：V1.2
- 编写人：孙丽华
- 审核人：周建国
- 日期：2026年3月8日
- 客户：浙江省人民医院

## 项目背景
浙江省人民医院（地址：杭州市上城区上塘路158号）计划建设智慧医疗信息系统，
预算总额为2,500万元，项目周期18个月。联系人：院信息科科长 陈志远，
电话：0571-87654321，邮箱：chenzy@zjhospital.cn

## 功能需求
1. 电子病历管理：支持结构化病历录入，日均处理量不低于5000份
2. 智能辅助诊断：基于AI的影像识别，准确率要求≥92%
3. 药品管理：对接国家药品监管平台，实时更新药品信息
4. 预约挂号：支持微信、支付宝、APP多渠道预约

## 非功能需求
- 系统可用性：99.9%
- 并发用户数：≥2000
- 响应时间：核心页面≤2秒
""",
    },
    {
        "filename": "竞品分析.md", "type": "md",
        "content": """# 智能文档处理产品竞品分析

## 分析人信息
- 分析师：林小雨
- 部门：产品研发部
- 日期：2026年3月12日
- 邮箱：linxy@company.com

## 竞品概览

### 产品A：百度文心一格
- 公司：百度在线网络技术（北京）有限公司
- 价格：企业版98,000元/年
- 市场份额：约15%

### 产品B：阿里云智能文档
- 公司：阿里云计算有限公司
- 价格：按量计费，约0.05元/页
- 市场份额：约22%

### 产品C：腾讯文档智能
- 公司：深圳市腾讯计算机系统有限公司
- 价格：企业版68,000元/年
- 市场份额：约18%

## 我方优势
1. 支持本地部署，数据不出企业
2. 价格优势：年费仅58,000元
3. 定制化能力强，支持私有化部署
""",
    },
]

XLSX_DOCS = [
    {
        "filename": "员工花名册.xlsx",
        "headers": ["工号", "姓名", "部门", "职位", "入职日期", "手机号", "邮箱", "身份证号"],
        "rows": [
            ["EMP001", "张三丰", "技术部", "高级工程师", "2023-03-15", "13811112222", "zhangsf@company.com", "110101199001011234"],
            ["EMP002", "李小龙", "市场部", "市场经理", "2022-07-01", "13922223333", "lixl@company.com", "310101199205152345"],
            ["EMP003", "王芳芳", "财务部", "财务主管", "2021-01-10", "13633334444", "wangff@company.com", "440101199108203456"],
            ["EMP004", "赵子龙", "技术部", "架构师", "2020-06-20", "13744445555", "zhaozl@company.com", "510101198812124567"],
            ["EMP005", "孙悟空", "运营部", "运营总监", "2024-02-28", "13855556666", "sunwk@company.com", "330101199503085678"],
        ],
    },
    {
        "filename": "销售数据统计.xlsx",
        "headers": ["月份", "客户名称", "产品", "金额(元)", "销售人员", "联系电话", "状态"],
        "rows": [
            ["2026-01", "北京华信科技", "企业版", "128000", "张伟", "13900001111", "已签约"],
            ["2026-01", "上海数据科技", "旗舰版", "258000", "李娜", "13900002222", "已签约"],
            ["2026-02", "广州云计算中心", "企业版", "128000", "王强", "13900003333", "已签约"],
            ["2026-02", "深圳智能科技", "定制版", "500000", "赵敏", "13900004444", "洽谈中"],
            ["2026-03", "成都天府数据", "企业版", "128000", "钱进", "13900005555", "已签约"],
        ],
    },
    {
        "filename": "项目进度表.xlsx",
        "headers": ["项目编号", "项目名称", "负责人", "开始日期", "截止日期", "预算(万元)", "进度", "状态"],
        "rows": [
            ["PRJ-001", "数据中台建设", "刘洋", "2026-01-15", "2026-06-30", "120", "45%", "进行中"],
            ["PRJ-002", "移动APP开发", "陈思", "2026-02-01", "2026-05-31", "80", "60%", "进行中"],
            ["PRJ-003", "安全审计系统", "周明", "2025-11-01", "2026-03-31", "95", "90%", "收尾中"],
            ["PRJ-004", "智慧园区平台", "吴刚", "2026-03-01", "2026-12-31", "350", "10%", "启动中"],
        ],
    },
    {
        "filename": "供应商信息表.xlsx",
        "headers": ["供应商名称", "联系人", "电话", "邮箱", "地址", "合作金额(万元)", "评级"],
        "rows": [
            ["华为技术有限公司", "张华", "0755-28780808", "zhanghua@huawei.com", "深圳市龙岗区坂田华为基地", "500", "A"],
            ["阿里云计算有限公司", "李云", "0571-85022088", "liyun@alibaba.com", "杭州市余杭区文一西路969号", "320", "A"],
            ["北京金山云信息技术", "王金", "010-62927777", "wangjin@ksyun.com", "北京市海淀区小营西路33号", "150", "B"],
            ["腾讯云计算有限公司", "赵腾", "0755-86013388", "zhaoteng@tencent.com", "深圳市南山区科技中一路", "280", "A"],
        ],
    },
]

# ============================================================
# 5个模板表格（含空字段，用于自动填写测试）
# ============================================================

TEMPLATES = [
    {
        "filename": "人员信息登记表.xlsx",
        "headers": ["姓名", "性别", "身份证号", "手机号", "邮箱", "部门", "职位", "入职日期"],
        "rows": [["", "", "", "", "", "", "", ""], ["", "", "", "", "", "", "", ""]],
    },
    {
        "filename": "项目立项申请表.xlsx",
        "headers": ["项目名称", "项目编号", "负责人", "联系电话", "预算金额", "开始日期", "结束日期", "客户名称"],
        "rows": [["", "", "", "", "", "", "", ""]],
    },
    {
        "filename": "合同信息登记表.xlsx",
        "headers": ["合同编号", "甲方名称", "甲方联系人", "甲方电话", "乙方名称", "合同金额", "签订日期", "到期日期"],
        "rows": [["", "", "", "", "", "", "", ""], ["", "", "", "", "", "", "", ""]],
    },
    {
        "filename": "客户信息采集表.xlsx",
        "headers": ["公司名称", "联系人", "职位", "电话", "邮箱", "地址", "合作意向", "备注"],
        "rows": [["", "", "", "", "", "", "", ""], ["", "", "", "", "", "", "", ""], ["", "", "", "", "", "", "", ""]],
    },
    {
        "filename": "费用报销单.xlsx",
        "headers": ["报销人", "部门", "日期", "费用类型", "金额", "事由", "审批人"],
        "rows": [["", "", "", "", "", "", ""], ["", "", "", "", "", "", ""]],
    },
]


# ============================================================
# 生成函数
# ============================================================

def generate_docx(data: dict):
    doc = DocxDocument()
    title_p = doc.add_heading(data.get("title", data["filename"]), level=1)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for para_text in data["content"]:
        for line in para_text.split("\n"):
            p = doc.add_paragraph(line)
            p.style.font.size = Pt(11)
    path = OUTPUT_DIR / data["filename"]
    doc.save(str(path))
    print(f"  [docx] {path.name}")
    return path


def generate_txt(data: dict):
    path = OUTPUT_DIR / data["filename"]
    path.write_text("\n\n".join(data["content"]), encoding="utf-8")
    print(f"  [txt]  {path.name}")
    return path


def generate_md(data: dict):
    path = OUTPUT_DIR / data["filename"]
    path.write_text(data["content"], encoding="utf-8")
    print(f"  [md]   {path.name}")
    return path


def generate_xlsx(data: dict):
    wb = Workbook()
    ws = wb.active
    ws.title = "数据"
    ws.append(data["headers"])
    header_fill = PatternFill(start_color="5B8DEF", end_color="5B8DEF", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True)
    for col_idx in range(1, len(data["headers"]) + 1):
        cell = ws.cell(row=1, column=col_idx)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center")
    for row in data["rows"]:
        ws.append(row)
    for i, _ in enumerate(data["headers"], 1):
        ws.column_dimensions[get_column_letter(i)].width = 18
    path = OUTPUT_DIR / data["filename"]
    wb.save(str(path))
    print(f"  [xlsx] {path.name}")
    return path


def generate_template(data: dict):
    wb = Workbook()
    ws = wb.active
    ws.title = "模板"
    ws.append(data["headers"])
    header_fill = PatternFill(start_color="FAAD14", end_color="FAAD14", fill_type="solid")
    header_font = Font(bold=True, size=11)
    for col_idx in range(1, len(data["headers"]) + 1):
        cell = ws.cell(row=1, column=col_idx)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center")
        ws.column_dimensions[get_column_letter(col_idx)].width = 20
    for row in data["rows"]:
        ws.append(row)
    path = TEMPLATE_DIR / data["filename"]
    wb.save(str(path))
    print(f"  [模板] {path.name}")
    return path


def main():
    print("=" * 50)
    print("生成测试文档和模板表格")
    print("=" * 50)

    print(f"\n输出目录: {OUTPUT_DIR}")
    print(f"模板目录: {TEMPLATE_DIR}\n")

    # 生成 docx
    print("--- docx 文档 ---")
    for d in DOCS_DATA:
        if d["type"] == "docx":
            generate_docx(d)

    # 生成 txt
    print("--- txt 文档 ---")
    for d in DOCS_DATA:
        if d["type"] == "txt":
            generate_txt(d)

    # 生成 md
    print("--- md 文档 ---")
    for d in MD_DOCS:
        generate_md(d)

    # 生成 xlsx
    print("--- xlsx 文档 ---")
    for d in XLSX_DOCS:
        generate_xlsx(d)

    # 生成模板
    print("--- 模板表格 ---")
    for t in TEMPLATES:
        generate_template(t)

    total_docs = len([d for d in DOCS_DATA]) + len(MD_DOCS) + len(XLSX_DOCS)
    print(f"\n完成! 共生成 {total_docs} 个测试文档 + {len(TEMPLATES)} 个模板表格")


if __name__ == "__main__":
    main()
