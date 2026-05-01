"""文档解析器 - 支持 docx/md/xlsx/txt/pdf/image 格式"""
from pathlib import Path
from datetime import datetime
from typing import Callable
from config import OCR_CONFIG
from logger import get_logger

log = get_logger("core.parser")


class ParserAdapter:
    """Small adapter wrapper for one document format family."""

    def __init__(self, suffixes: set[str], parse_func: Callable[[Path], str]):
        self.suffixes = suffixes
        self.parse_func = parse_func

    def parse(self, path: Path) -> str:
        return self.parse_func(path)


class DocumentParser:
    IMAGE_TYPES = {".png", ".jpg", ".jpeg", ".bmp"}
    SUPPORTED_TYPES = {".docx", ".md", ".xlsx", ".txt", ".pdf"} | IMAGE_TYPES
    _CACHE: dict[tuple[str, float, int], dict] = {}
    _ADAPTERS: dict[str, ParserAdapter] = {}

    @staticmethod
    def parse(file_path: str) -> dict:
        """解析文档，返回 {"text": str, "file_type": str, "metadata": dict}"""
        path = Path(file_path)
        suffix = path.suffix.lower()
        if suffix not in DocumentParser.SUPPORTED_TYPES:
            raise ValueError(f"不支持的文件格式: {suffix}")

        stat = path.stat() if path.exists() else None
        file_size = stat.st_size if stat else 0
        mtime = stat.st_mtime if stat else 0.0
        cache_key = (str(path.resolve()), mtime, file_size)
        cached = DocumentParser._CACHE.get(cache_key)
        if cached:
            return {
                "text": cached["text"],
                "file_type": cached["file_type"],
                "metadata": {
                    **cached["metadata"],
                    "parsed_at": datetime.now().isoformat(),
                    "cache_hit": True,
                },
            }

        metadata = {
            "filename": path.name,
            "file_size": file_size,
            "parsed_at": datetime.now().isoformat(),
            "cache_hit": False,
        }

        try:
            adapter = DocumentParser._get_adapter(suffix)
            text = adapter.parse(path)
        except Exception as e:
            log.error(f"解析文件失败 {path.name}: {e}")
            raise

        log.info(f"解析完成: {path.name}, {len(text)}字符")
        result = {"text": text, "file_type": suffix.lstrip("."), "metadata": metadata}
        DocumentParser._CACHE[cache_key] = result
        if len(DocumentParser._CACHE) > 64:
            oldest_key = next(iter(DocumentParser._CACHE))
            DocumentParser._CACHE.pop(oldest_key, None)
        return result

    @classmethod
    def register_adapter(cls, adapter: ParserAdapter):
        for suffix in adapter.suffixes:
            cls._ADAPTERS[suffix] = adapter
        cls.SUPPORTED_TYPES = set(cls._ADAPTERS)

    @classmethod
    def _get_adapter(cls, suffix: str) -> ParserAdapter:
        if not cls._ADAPTERS:
            cls._register_default_adapters()
        adapter = cls._ADAPTERS.get(suffix)
        if not adapter:
            raise ValueError(f"不支持的文件格式: {suffix}")
        return adapter

    @classmethod
    def _register_default_adapters(cls):
        cls.register_adapter(ParserAdapter({".txt"}, lambda path: cls._parse_txt(path)))
        cls.register_adapter(ParserAdapter({".md"}, lambda path: cls._parse_md(path)))
        cls.register_adapter(ParserAdapter({".docx"}, lambda path: cls._parse_docx(path)))
        cls.register_adapter(ParserAdapter({".xlsx"}, lambda path: cls._parse_xlsx(path)))
        cls.register_adapter(ParserAdapter({".pdf"}, lambda path: cls._parse_pdf(path)))
        cls.register_adapter(ParserAdapter(cls.IMAGE_TYPES, lambda path: cls._parse_image(path)))

    @staticmethod
    def _parse_md(path: Path) -> str:
        return path.read_text(encoding="utf-8")

    @staticmethod
    def _parse_txt(path: Path) -> str:
        return path.read_text(encoding="utf-8")

    @staticmethod
    def _parse_docx(path: Path) -> str:
        from docx import Document as DocxDocument
        doc = DocxDocument(str(path))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)

    @staticmethod
    def _parse_xlsx(path: Path) -> str:
        from openpyxl import load_workbook
        wb = load_workbook(str(path), read_only=True, data_only=True)
        parts = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            parts.append(f"[工作表: {sheet_name}]")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None]
                if cells:
                    parts.append(" | ".join(cells))
        wb.close()
        return "\n".join(parts)

    @staticmethod
    def _parse_pdf(path: Path) -> str:
        try:
            import fitz  # PyMuPDF
            text_parts = []
            with fitz.open(str(path)) as doc:
                for page in doc:
                    page_text = page.get_text()
                    if page_text.strip():
                        text_parts.append(page_text)
            return "\n".join(text_parts)
        except ImportError:
            log.warning("未安装PyMuPDF，跳过PDF解析。请运行: pip install PyMuPDF")
            return "[PDF解析需要安装 PyMuPDF]"

    @staticmethod
    def _parse_image(path: Path) -> str:
        try:
            import pytesseract
            from PIL import Image
        except ImportError:
            log.warning("未安装pytesseract或Pillow，跳过图片OCR。请运行: pip install pytesseract Pillow")
            return "[图片OCR需要安装 pytesseract 和 Pillow]"

        tesseract_cmd = OCR_CONFIG.get("tesseract_cmd")
        if tesseract_cmd:
            pytesseract.pytesseract.tesseract_cmd = tesseract_cmd

        try:
            with Image.open(path) as image:
                return pytesseract.image_to_string(image, lang=OCR_CONFIG.get("lang", "chi_sim+eng")).strip()
        except pytesseract.TesseractNotFoundError as e:
            msg = f"Tesseract-OCR未找到，请检查OCR_CONFIG.tesseract_cmd配置: {e}"
            log.error(msg)
            raise RuntimeError(msg) from e
