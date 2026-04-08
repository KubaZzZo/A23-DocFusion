"""模板表格填写器"""
import json
import shutil
from pathlib import Path
from datetime import datetime
from openpyxl import load_workbook
from docx import Document as DocxDocument
from core.semantic_matcher import SemanticMatcher
from config import OUTPUT_DIR


class TemplateFiller:
    def __init__(self, provider: str = None):
        self.matcher = SemanticMatcher(provider)

    async def analyze_template(self, file_path: str) -> dict:
        """分析模板，识别需要填写的字段"""
        path = Path(file_path)
        suffix = path.suffix.lower()
        if suffix == ".xlsx":
            return self._analyze_xlsx(path)
        elif suffix == ".docx":
            return self._analyze_docx(path)
        else:
            raise ValueError(f"不支持的模板格式: {suffix}")

    def _analyze_xlsx(self, path: Path) -> dict:
        """分析xlsx模板，找出表头和空单元格"""
        wb = load_workbook(str(path))
        fields = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            headers = []
            # 第一行作为表头
            for cell in next(ws.iter_rows(min_row=1, max_row=1, values_only=False), []):
                if cell.value:
                    headers.append({"col": cell.column, "name": str(cell.value), "sheet": sheet_name})

            # 扫描空单元格
            for row in ws.iter_rows(min_row=2, values_only=False):
                for cell in row:
                    if cell.value is None or str(cell.value).strip() == "":
                        # 找到对应表头
                        header = next((h for h in headers if h["col"] == cell.column), None)
                        if header:
                            fields.append({
                                "field_name": header["name"],
                                "sheet": sheet_name,
                                "row": cell.row,
                                "col": cell.column,
                            })
        wb.close()
        return {"fields": fields, "field_names": list(set(f["field_name"] for f in fields))}

    def _analyze_docx(self, path: Path) -> dict:
        """分析docx模板中的表格，找出空单元格"""
        doc = DocxDocument(str(path))
        fields = []
        for t_idx, table in enumerate(doc.tables):
            headers = [cell.text.strip() for cell in table.rows[0].cells] if table.rows else []
            for r_idx, row in enumerate(table.rows[1:], start=1):
                for c_idx, cell in enumerate(row.cells):
                    if not cell.text.strip() and c_idx < len(headers) and headers[c_idx]:
                        fields.append({
                            "field_name": headers[c_idx],
                            "table_index": t_idx,
                            "row": r_idx,
                            "col": c_idx,
                        })
        return {"fields": fields, "field_names": list(set(f["field_name"] for f in fields))}

    async def fill(self, template_path: str, entities: list[dict]) -> dict:
        """自动填写模板"""
        # 1. 分析模板
        analysis = await self.analyze_template(template_path)
        field_names = analysis["field_names"]

        if not field_names:
            return {"success": False, "message": "模板中未找到需要填写的字段"}

        # 2. 语义匹配
        match_result = await self.matcher.match(field_names, entities)
        matches = match_result.get("matches", [])

        # 3. 构建填写映射
        fill_map = {}
        for m in matches:
            fill_map[m["field"]] = m["value"]

        # 4. 执行填写
        path = Path(template_path)
        suffix = path.suffix.lower()
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_name = f"{path.stem}_filled_{timestamp}{suffix}"
        output_path = OUTPUT_DIR / output_name
        shutil.copyfile(template_path, output_path)
        output_path.chmod(0o666)

        if suffix == ".xlsx":
            self._fill_xlsx(str(output_path), analysis["fields"], fill_map)
        elif suffix == ".docx":
            self._fill_docx(str(output_path), analysis["fields"], fill_map)

        filled_count = sum(1 for f in analysis["fields"] if f["field_name"] in fill_map)
        total_count = len(analysis["fields"])

        return {
            "success": True,
            "output_path": str(output_path),
            "filled": filled_count,
            "total": total_count,
            "accuracy": filled_count / total_count if total_count > 0 else 0,
            "unmatched": match_result.get("unmatched_fields", []),
        }

    def _fill_xlsx(self, file_path: str, fields: list[dict], fill_map: dict):
        wb = load_workbook(file_path)
        for f in fields:
            value = fill_map.get(f["field_name"])
            if value:
                ws = wb[f["sheet"]]
                ws.cell(row=f["row"], column=f["col"], value=value)
        wb.save(file_path)
        wb.close()

    def _fill_docx(self, file_path: str, fields: list[dict], fill_map: dict):
        doc = DocxDocument(file_path)
        for f in fields:
            value = fill_map.get(f["field_name"])
            if value:
                table = doc.tables[f["table_index"]]
                table.rows[f["row"]].cells[f["col"]].text = value
        doc.save(file_path)
