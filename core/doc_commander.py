"""文档智能操作 - 自然语言指令解析与执行"""
import json
import shutil
from pathlib import Path
from datetime import datetime
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from llm import get_llm
from llm.base import strip_json_code_fence
from config import DATA_DIR

BACKUP_DIR = DATA_DIR / "backups"
BACKUP_DIR.mkdir(exist_ok=True)

COMMAND_PARSE_PROMPT = """你是一个文档操作指令解析器。用户会用自然语言描述对文档的操作需求，你需要将其转换为结构化的操作指令JSON。

支持的操作类型：
1. format - 格式调整（bold, italic, underline, font_size, font_name, color, alignment）
2. edit - 内容编辑（insert, delete, replace）
3. extract - 内容提取（extract_text, extract_tables, extract_headings）
4. find_replace - 查找替换
5. structure - 结构操作（add_heading, add_paragraph, add_table）

请输出JSON格式：
{
  "action": "操作类型",
  "target": "操作目标(paragraph/table/heading/all)",
  "params": {具体参数},
  "description": "操作描述"
}

示例：
用户："把第二段加粗" → {"action":"format","target":"paragraph","params":{"index":1,"bold":true},"description":"将第2段加粗"}
用户："查找所有'公司'替换为'企业'" → {"action":"find_replace","target":"all","params":{"find":"公司","replace":"企业"},"description":"全文替换"}
用户："提取所有表格" → {"action":"extract","target":"tables","params":{},"description":"提取所有表格内容"}
"""


class DocCommander:
    """文档智能操作执行器"""

    def __init__(self, provider: str = None):
        self.llm = get_llm(provider)

    async def parse_command(self, user_input: str, doc_info: str = "") -> dict:
        """解析自然语言指令为操作JSON"""
        messages = [
            {"role": "system", "content": COMMAND_PARSE_PROMPT},
            {"role": "user", "content": f"文档信息：{doc_info}\n\n用户指令：{user_input}"},
        ]
        result = await self.llm.chat(messages)
        try:
            cleaned = strip_json_code_fence(result)
            parsed = json.loads(cleaned)
            return self._normalize_command(user_input, parsed)
        except json.JSONDecodeError:
            return {"error": "指令解析失败", "raw": result}

    @staticmethod
    def _normalize_command(user_input: str, parsed: dict) -> dict:
        """对 LLM 输出做轻量纠偏，避免常见文档结构歧义。"""
        if not isinstance(parsed, dict):
            return parsed

        action = parsed.get("action")
        target = parsed.get("target")
        params = parsed.get("params", {})
        description = parsed.get("description", "")
        text = f"{user_input} {description}"

        if (
            action == "format"
            and target == "paragraph"
            and isinstance(params, dict)
            and "index" in params
            and "行" in text
        ):
            parsed["target"] = "table_row"

        return parsed

    def execute(self, doc_path: str, command: dict) -> dict:
        """执行操作指令（自动备份原文件）"""
        if Path(doc_path).suffix.lower() != ".docx":
            return {"success": False, "message": "文档智能操作目前仅支持 .docx 格式"}

        action = command.get("action")
        handlers = {
            "format": self._handle_format,
            "edit": self._handle_edit,
            "find_replace": self._handle_find_replace,
            "extract": self._handle_extract,
            "structure": self._handle_structure,
        }
        handler = handlers.get(action)
        if not handler:
            return {"success": False, "message": f"不支持的操作: {action}"}

        # extract 是只读操作，不需要备份
        if action != "extract":
            backup_path = self._backup(doc_path)
        else:
            backup_path = None

        try:
            params = dict(command.get("params", {}))
            if "target" in command and "target" not in params:
                params["target"] = command["target"]
            result = handler(doc_path, params)
            if backup_path:
                result["backup_path"] = str(backup_path)
            return result
        except Exception as e:
            # 操作失败时自动恢复
            if backup_path and Path(backup_path).exists():
                shutil.copyfile(backup_path, doc_path)
                Path(doc_path).chmod(0o666)
            return {"success": False, "message": str(e)}

    @staticmethod
    def _backup(doc_path: str) -> Path:
        """备份文件，返回备份路径"""
        src = Path(doc_path)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        backup_path = BACKUP_DIR / f"{src.stem}_{timestamp}{src.suffix}"
        shutil.copyfile(doc_path, backup_path)
        backup_path.chmod(0o666)
        return backup_path

    def _handle_format(self, doc_path: str, params: dict) -> dict:
        doc = DocxDocument(doc_path)
        target = params.get("target", "paragraph")

        if target == "table_row":
            row_idx = params.get("index", 0)
            tables = doc.tables
            if not tables:
                return {"success": False, "message": "当前文档中没有表格，无法按“行”执行格式操作"}
            first_table = tables[0]
            if row_idx >= len(first_table.rows):
                return {"success": False, "message": f"表格行索引{row_idx}超出范围"}

            for cell in first_table.rows[row_idx].cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if "bold" in params:
                            run.bold = params["bold"]
                        if "italic" in params:
                            run.italic = params["italic"]
                        if "underline" in params:
                            run.underline = params["underline"]
                        if "font_size" in params:
                            run.font.size = Pt(params["font_size"])
                        if "font_name" in params:
                            run.font.name = params["font_name"]
                        if "color" in params:
                            r, g, b = params["color"]
                            run.font.color.rgb = RGBColor(r, g, b)
                if "alignment" in params:
                    align_map = {
                        "left": WD_ALIGN_PARAGRAPH.LEFT,
                        "center": WD_ALIGN_PARAGRAPH.CENTER,
                        "right": WD_ALIGN_PARAGRAPH.RIGHT,
                        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
                    }
                    for para in cell.paragraphs:
                        para.alignment = align_map.get(params["alignment"], WD_ALIGN_PARAGRAPH.LEFT)

            doc.save(doc_path)
            return {"success": True, "message": f"已完成第{row_idx + 1}行格式调整"}

        idx = params.get("index", 0)
        if idx >= len(doc.paragraphs):
            return {"success": False, "message": f"段落索引{idx}超出范围"}

        para = doc.paragraphs[idx]
        for run in para.runs:
            if "bold" in params:
                run.bold = params["bold"]
            if "italic" in params:
                run.italic = params["italic"]
            if "underline" in params:
                run.underline = params["underline"]
            if "font_size" in params:
                run.font.size = Pt(params["font_size"])
            if "font_name" in params:
                run.font.name = params["font_name"]
            if "color" in params:
                r, g, b = params["color"]
                run.font.color.rgb = RGBColor(r, g, b)

        if "alignment" in params:
            align_map = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                         "right": WD_ALIGN_PARAGRAPH.RIGHT, "justify": WD_ALIGN_PARAGRAPH.JUSTIFY}
            para.alignment = align_map.get(params["alignment"], WD_ALIGN_PARAGRAPH.LEFT)

        doc.save(doc_path)
        return {"success": True, "message": "格式调整完成"}

    def _handle_edit(self, doc_path: str, params: dict) -> dict:
        doc = DocxDocument(doc_path)
        op = params.get("operation", "replace")

        if op == "replace" and "index" in params:
            idx = params["index"]
            if idx < len(doc.paragraphs):
                doc.paragraphs[idx].text = params.get("text", "")
        elif op == "insert":
            doc.add_paragraph(params.get("text", ""))
        elif op == "delete" and "index" in params:
            idx = params["index"]
            if idx < len(doc.paragraphs):
                p = doc.paragraphs[idx]._element
                p.getparent().remove(p)

        doc.save(doc_path)
        return {"success": True, "message": "编辑完成"}

    def _handle_find_replace(self, doc_path: str, params: dict) -> dict:
        doc = DocxDocument(doc_path)
        find_text = params.get("find", "")
        replace_text = params.get("replace", "")
        count = 0
        if not find_text:
            return {"success": False, "message": "查找内容不能为空"}

        for para in doc.paragraphs:
            if find_text not in para.text:
                continue
            count += self._replace_in_paragraph_runs(para, find_text, replace_text)

        doc.save(doc_path)
        return {"success": True, "message": f"替换完成，共替换{count}处"}

    @staticmethod
    def _replace_in_paragraph_runs(para, find_text: str, replace_text: str) -> int:
        """在 run 层面执行替换，尽量保留未命中内容的原有格式。"""
        matches = []
        full_text_parts = []
        char_map = []

        for run_idx, run in enumerate(para.runs):
            for char_idx, ch in enumerate(run.text):
                full_text_parts.append(ch)
                char_map.append((run_idx, char_idx))

        full_text = "".join(full_text_parts)
        start = 0
        while True:
            index = full_text.find(find_text, start)
            if index == -1:
                break
            matches.append((index, index + len(find_text)))
            start = index + len(find_text)

        for start_idx, end_idx in reversed(matches):
            start_run_idx, start_char_idx = char_map[start_idx]
            end_run_idx, end_char_idx = char_map[end_idx - 1]
            start_run = para.runs[start_run_idx]

            if start_run_idx == end_run_idx:
                text = start_run.text
                start_run.text = text[:start_char_idx] + replace_text + text[end_char_idx + 1:]
                continue

            start_run.text = start_run.text[:start_char_idx] + replace_text
            for run_idx in range(start_run_idx + 1, end_run_idx):
                para.runs[run_idx].text = ""
            end_run = para.runs[end_run_idx]
            end_run.text = end_run.text[end_char_idx + 1:]

        return len(matches)

    def _handle_extract(self, doc_path: str, params: dict) -> dict:
        doc = DocxDocument(doc_path)
        target = params.get("target", "text")

        if target == "tables":
            tables = []
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    rows.append([cell.text for cell in row.cells])
                tables.append(rows)
            return {"success": True, "data": tables}
        elif target == "headings":
            headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
            return {"success": True, "data": headings}
        else:
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            return {"success": True, "data": text}

    def _handle_structure(self, doc_path: str, params: dict) -> dict:
        doc = DocxDocument(doc_path)
        op = params.get("operation", "add_paragraph")

        if op == "add_heading":
            doc.add_heading(params.get("text", ""), level=params.get("level", 1))
        elif op == "add_paragraph":
            doc.add_paragraph(params.get("text", ""))

        doc.save(doc_path)
        return {"success": True, "message": "结构操作完成"}
