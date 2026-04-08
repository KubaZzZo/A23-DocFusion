"""将爬取的文章生成多种格式的测试文档"""
from pathlib import Path
from datetime import datetime
from docx import Document as DocxDocument
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from config import DATA_DIR

CRAWLED_DIR = DATA_DIR / "crawled"
CRAWLED_DIR.mkdir(exist_ok=True)


class DocGenerator:
    @staticmethod
    def generate_docx(article: dict) -> str:
        """生成带标题、正文、信息表格的docx"""
        doc = DocxDocument()

        # 标题
        title_p = doc.add_heading(article.get("title", "无标题"), level=1)
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 信息表格
        table = doc.add_table(rows=4, cols=2, style="Light Grid Accent 1")
        fields = [
            ("来源", article.get("source", "")),
            ("作者", article.get("author", "")),
            ("发布日期", article.get("publish_date", "")),
            ("链接", article.get("url", "")),
        ]
        for i, (k, v) in enumerate(fields):
            table.cell(i, 0).text = k
            table.cell(i, 1).text = v

        doc.add_paragraph("")

        # 正文
        content = article.get("content", "")
        for para in content.split("\n"):
            para = para.strip()
            if para:
                p = doc.add_paragraph(para)
                p.style.font.size = Pt(11)

        safe_title = _safe_filename(article.get("title", "article"))
        path = CRAWLED_DIR / f"{safe_title}.docx"
        doc.save(str(path))
        return str(path)

    @staticmethod
    def generate_xlsx(articles: list[dict]) -> str:
        """多篇文章汇总到xlsx表格"""
        wb = Workbook()
        ws = wb.active
        ws.title = "爬取文章汇总"

        headers = ["序号", "标题", "来源", "作者", "发布日期", "内容摘要", "链接"]
        ws.append(headers)

        # 设置表头样式
        from openpyxl.styles import Font, PatternFill, Alignment
        header_fill = PatternFill(start_color="5B8DEF", end_color="5B8DEF", fill_type="solid")
        header_font = Font(color="FFFFFF", bold=True, size=11)
        for col_idx, _ in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col_idx)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center")

        for i, a in enumerate(articles, 1):
            content = a.get("content", "")
            summary = content[:200] + "..." if len(content) > 200 else content
            ws.append([
                i,
                a.get("title", ""),
                a.get("source", ""),
                a.get("author", ""),
                a.get("publish_date", ""),
                summary,
                a.get("url", ""),
            ])

        # 调整列宽
        ws.column_dimensions["A"].width = 6
        ws.column_dimensions["B"].width = 40
        ws.column_dimensions["C"].width = 14
        ws.column_dimensions["D"].width = 14
        ws.column_dimensions["E"].width = 16
        ws.column_dimensions["F"].width = 60
        ws.column_dimensions["G"].width = 40

        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        path = CRAWLED_DIR / f"articles_{ts}.xlsx"
        wb.save(str(path))
        return str(path)

    @staticmethod
    def generate_txt(article: dict) -> str:
        """纯文本格式"""
        lines = [
            f"标题: {article.get('title', '')}",
            f"来源: {article.get('source', '')}",
            f"作者: {article.get('author', '')}",
            f"日期: {article.get('publish_date', '')}",
            f"链接: {article.get('url', '')}",
            "",
            "=" * 50,
            "",
            article.get("content", ""),
        ]
        safe_title = _safe_filename(article.get("title", "article"))
        path = CRAWLED_DIR / f"{safe_title}.txt"
        path.write_text("\n".join(lines), encoding="utf-8")
        return str(path)

    @staticmethod
    def generate_md(article: dict) -> str:
        """Markdown格式"""
        lines = [
            f"# {article.get('title', '')}",
            "",
            f"> 来源: {article.get('source', '')} | 作者: {article.get('author', '')} | 日期: {article.get('publish_date', '')}",
            "",
            f"[原文链接]({article.get('url', '')})",
            "",
            "---",
            "",
            article.get("content", ""),
        ]
        safe_title = _safe_filename(article.get("title", "article"))
        path = CRAWLED_DIR / f"{safe_title}.md"
        path.write_text("\n".join(lines), encoding="utf-8")
        return str(path)

    @staticmethod
    def generate_all(articles: list[dict]) -> dict:
        """批量生成所有格式，返回生成的文件路径"""
        gen = DocGenerator()
        paths = {"docx": [], "txt": [], "md": []}
        for a in articles:
            if a.get("content"):
                paths["docx"].append(gen.generate_docx(a))
                paths["txt"].append(gen.generate_txt(a))
                paths["md"].append(gen.generate_md(a))
        if articles:
            paths["xlsx"] = [gen.generate_xlsx(articles)]
        return paths


def _safe_filename(name: str, max_len: int = 50) -> str:
    """清理文件名中的非法字符"""
    import re
    name = re.sub(r'[\\/:*?"<>|\n\r]', '_', name)
    return name[:max_len].strip('_. ') or "article"
